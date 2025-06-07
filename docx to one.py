import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_name, bookmark_id):
    """ 向段落添加书签 """
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    paragraph._p.append(bookmark_start)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    bookmark_end.set(qn('w:name'), bookmark_name)
    paragraph._p.append(bookmark_end)

def add_hyperlink(paragraph, text, anchor):
    """ 添加超链接到段落 """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def read_file_content(file_path):
    """ 读取不同格式文件内容 """
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

def merge_files(input_files, output_path):
    """ 整合文件并生成文档 """
    doc = Document()
    files_data = []
    
    # 预处理所有文件
    for idx, file_path in enumerate(input_files):
        file_name = os.path.basename(file_path)
        title = os.path.splitext(file_name)[0].replace('_', ' ').replace('-', ' ')
        bookmark = f'BOOKMARK_{idx}'
        content = read_file_content(file_path)
        files_data.append((title, bookmark, content))
    
    # 添加目录
    toc_para = doc.add_paragraph()
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_para.add_run('目录')
    run.font.size = Pt(24)
    
    for title, bookmark, _ in files_data:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(para, title, bookmark)
        for run in para.runs:
            run.font.size = Pt(13.5)
    
    doc.add_page_break()
    
    # 添加正文内容
    for idx, (title, bookmark, content) in enumerate(files_data):
        # 添加标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(24)
        add_bookmark(title_para, bookmark, idx)
        
        # 添加正文
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(13.5)
        
        # 添加分页符（最后一个文件不添加）
        if idx < len(files_data) - 1:
            doc.add_page_break()
    
    doc.save(output_path)

if __name__ == "__main__":
    # 使用示例
    input_files = ['file1.txt', 'file2.docx']  # 替换为实际文件路径
    output_path = 'merged_document.docx'
    merge_files(input_files, output_path)